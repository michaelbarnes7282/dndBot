import os
from openai import OpenAI
from docx import Document
from dotenv import load_dotenv

load_dotenv()

OPENAI_KEY = os.getenv("OPEN_API_KEY")
client = OpenAI(
    api_key=OPENAI_KEY
)


def transcribe_audio(audio_file_path):
    audio_file = open(audio_file_path, "rb")
    transcription = client.audio.transcriptions.create(
        model="whisper-1",
        file=audio_file,
        response_format="text"
    )
    return transcription


def meeting_minutes(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)
    sentiment = sentiment_analysis(transcription)
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }


def abstract_summary_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "As a masterful AI scribe with expertise in language comprehension and summarization, your task is to chronicle the unfolding adventures of a D&D session. Upon reading the provided narrative and dialogues, distill them into a concise abstract paragraph. Your focus should be on capturing the pivotal plot developments, significant character actions, and critical decisions made by the players. Aim to provide a coherent and readable summary that encapsulates the essence of the session, enabling one to grasp the main events and themes without needing to delve into the full depth of the role-playing experience. Please exclude any extraneous details or tangents, retaining only what is crucial to understanding the story's progression and the players' interactions."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content


def key_points_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "As a sagacious AI Chronicler, skilled in the art of discerning and distilling the essence of tales and lore, your quest is to dissect the chronicles of a D&D session. From the narrated adventures and dialogues, identify and catalog the most pivotal points and themes. These should encapsulate the key events, significant character decisions, and major plot twists – the crux of the session's story. Your aim is to forge a list that, when perused, offers anyone a swift and clear understanding of the session's narrative, the adventurers' journey, and the crucial moments that shaped their path. Focus on extracting the core elements that define the essence of the session's tale."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content


def action_item_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "As an AI Archivist, adept in the mystical art of interpreting the spoken word and uncovering hidden directives, your mission is to scrutinize the narratives and dialogues of a D&D session. Examine the intricacies of the adventurers' conversations and unearth any quests, objectives, or tasks that were agreed upon or noted as essential. These could be specific missions entrusted to individual heroes, or collective endeavors that the party has vowed to undertake. Enumerate these action items with clarity and precision, so that the adventurers may have a clear record of their commitments and future undertakings in their ongoing saga."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )

    return response.choices[0].message.content


def sentiment_analysis(transcription):
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "As an AI Seer with mastery in the arcane arts of language and emotion divination, your quest is to delve into the essence of the spoken and written words from a D&D session. Scrutinize the fabric of the narrative, the emotions woven by the language, and the context in which the words are cast. Discern the overarching sentiment of the tale – be it valiantly positive, ominously negative, or a balanced neutral. Provide insights into your analysis, shedding light on the emotional undertones and the mood that pervades the adventurers' journey. Your insights will illuminate the emotional landscape of the session, revealing the heart of the story as it unfolds."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )

    return response.choices[0].message.content


def save_as_docx(transcription, filename):
    minutes = meeting_minutes(transcription)
    doc = Document()
    for key, value in minutes.items():
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)
